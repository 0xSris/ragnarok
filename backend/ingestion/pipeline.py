"""
Ingestion pipeline — handles PDF, DOCX, TXT, images (OCR), and audio (Whisper).
Extracts text, chunks it, returns structured chunks for vector indexing.
"""
import io
import json
import logging
import re
import uuid
from pathlib import Path
from typing import List, Dict, Any, Tuple

logger = logging.getLogger(__name__)


# ─── Text Chunking ────────────────────────────────────────────────────────────

def chunk_text(
    text: str,
    chunk_size: int = 512,
    overlap: int = 64,
    source: str = "",
    doc_type: str = "",
    filename: str = "",
    page: int = 0,
) -> List[Dict[str, Any]]:
    """Split text into overlapping chunks with metadata."""
    words = text.split()
    chunks = []
    i = 0
    while i < len(words):
        window = words[i : i + chunk_size]
        chunk_text = " ".join(window).strip()
        if chunk_text:
            chunks.append({
                "text": chunk_text,
                "source": source,
                "doc_type": doc_type,
                "filename": filename,
                "page": page,
                "word_start": i,
                "word_end": i + len(window),
            })
        i += chunk_size - overlap
    return chunks


# ─── PDF ──────────────────────────────────────────────────────────────────────

def extract_pdf(file_path: Path) -> Tuple[List[Dict], int, int]:
    """Extract text from PDF. Falls back to OCR if text layer is empty."""
    import PyPDF2
    from pdf2image import convert_from_path
    import pytesseract

    chunks = []
    total_pages = 0
    total_words = 0
    filename = file_path.name

    try:
        with open(file_path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            total_pages = len(reader.pages)

            for page_num, page in enumerate(reader.pages):
                text = page.extract_text() or ""
                text = re.sub(r"\s+", " ", text).strip()

                # If text is too short, OCR the page
                if len(text) < 50:
                    try:
                        images = convert_from_path(
                            str(file_path),
                            first_page=page_num + 1,
                            last_page=page_num + 1,
                            dpi=200,
                        )
                        if images:
                            text = pytesseract.image_to_string(images[0])
                            text = re.sub(r"\s+", " ", text).strip()
                    except Exception as e:
                        logger.warning(f"OCR fallback failed page {page_num}: {e}")

                if text:
                    total_words += len(text.split())
                    page_chunks = chunk_text(
                        text,
                        source=str(file_path),
                        doc_type="pdf",
                        filename=filename,
                        page=page_num + 1,
                    )
                    chunks.extend(page_chunks)

    except Exception as e:
        logger.error(f"PDF extraction error: {e}")
        raise

    return chunks, total_pages, total_words


# ─── DOCX ─────────────────────────────────────────────────────────────────────

def extract_docx(file_path: Path) -> Tuple[List[Dict], int, int]:
    from docx import Document

    doc = Document(str(file_path))
    full_text = []
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text.strip())

    # Tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                full_text.append(row_text)

    text = "\n".join(full_text)
    total_words = len(text.split())
    chunks = chunk_text(
        text,
        source=str(file_path),
        doc_type="docx",
        filename=file_path.name,
        page=1,
    )
    return chunks, 1, total_words


# ─── TXT / MD ─────────────────────────────────────────────────────────────────

def extract_txt(file_path: Path) -> Tuple[List[Dict], int, int]:
    text = file_path.read_text(errors="replace")
    total_words = len(text.split())
    chunks = chunk_text(
        text,
        source=str(file_path),
        doc_type="txt",
        filename=file_path.name,
    )
    return chunks, 1, total_words


# ─── Images (OCR) ─────────────────────────────────────────────────────────────

def extract_image(file_path: Path) -> Tuple[List[Dict], int, int]:
    import pytesseract
    from PIL import Image

    img = Image.open(str(file_path))
    text = pytesseract.image_to_string(img)
    text = re.sub(r"\s+", " ", text).strip()
    total_words = len(text.split())
    chunks = chunk_text(
        text,
        source=str(file_path),
        doc_type="image",
        filename=file_path.name,
    )
    return chunks, 1, total_words


# ─── Audio (Whisper) ──────────────────────────────────────────────────────────

def extract_audio(file_path: Path) -> Tuple[List[Dict], int, int]:
    import whisper
    from backend.core.config import settings

    model = whisper.load_model(settings.WHISPER_MODEL)
    result = model.transcribe(str(file_path))
    text = result.get("text", "").strip()
    total_words = len(text.split())
    chunks = chunk_text(
        text,
        source=str(file_path),
        doc_type="audio",
        filename=file_path.name,
    )
    return chunks, 1, total_words


# ─── Dispatcher ───────────────────────────────────────────────────────────────

EXTRACTORS = {
    ".pdf": extract_pdf,
    ".docx": extract_docx,
    ".doc": extract_docx,
    ".txt": extract_txt,
    ".md": extract_txt,
    ".png": extract_image,
    ".jpg": extract_image,
    ".jpeg": extract_image,
    ".tiff": extract_image,
    ".bmp": extract_image,
    ".webp": extract_image,
    ".mp3": extract_audio,
    ".wav": extract_audio,
    ".m4a": extract_audio,
    ".ogg": extract_audio,
    ".flac": extract_audio,
    ".mp4": extract_audio,
}


def ingest_file(file_path: Path) -> Dict[str, Any]:
    """
    Main ingestion entry point.
    Returns: {chunks, page_count, word_count, file_type}
    """
    suffix = file_path.suffix.lower()
    extractor = EXTRACTORS.get(suffix)
    if not extractor:
        raise ValueError(f"Unsupported file type: {suffix}")

    chunks, pages, words = extractor(file_path)
    logger.info(
        f"Ingested {file_path.name}: {len(chunks)} chunks, "
        f"{pages} pages, {words} words"
    )
    return {
        "chunks": chunks,
        "page_count": pages,
        "word_count": words,
        "file_type": suffix.lstrip("."),
        "chunk_count": len(chunks),
    }
